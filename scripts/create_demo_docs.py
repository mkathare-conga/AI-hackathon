"""
Demo Document Generator
=======================
Generates realistic multi-page DOCX contract documents for any demo account.

Usage
-----
  python scripts/create_demo_docs.py --help

  # Generate all four document types for a new account
  python scripts/create_demo_docs.py \\
      --account "Acme Corp" \\
      --product "Conga Revenue Suite" \\
      --base-price 8000 \\
      --quantity 100 \\
      --uplift 5 \\
      --amendment-uplift 8 \\
      --term-start 2024-01-01

  # Generate only MSA and amendment
  python scripts/create_demo_docs.py \\
      --account "Contoso Ltd" \\
      --product "Conga CLM Pro" \\
      --base-price 12000 \\
      --quantity 200 \\
      --uplift 4 \\
      --amendment-uplift 9 \\
      --types msa amendment

  # Specify a custom output directory
  python scripts/create_demo_docs.py \\
      --account "Northstar Inc" \\
      --product "Conga Sign Enterprise" \\
      --base-price 5000 \\
      --quantity 50 \\
      --uplift 3 \\
      --output-dir demo/northstar
"""
from __future__ import annotations

import argparse
import re
import sys
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import NamedTuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


# ─── Config ──────────────────────────────────────────────────────────────────

class DocConfig(NamedTuple):
    account: str          # Customer company name
    supplier: str         # Supplier name (default: Conga)
    product: str          # Product / subscription name
    base_price: float     # Per-unit monthly price
    quantity: int         # Number of units/seats
    currency: str         # USD / EUR / GBP
    uplift: float         # MSA uplift % (e.g. 5.0)
    amendment_uplift: float  # Amendment override uplift % (e.g. 8.0)
    term_start: date
    term_end: date
    renewal_date: date    # Usually term_end + 1 day
    notice_days: int      # Notice window days (default 30)
    slug: str             # Filename prefix derived from account name


def make_config(args) -> DocConfig:
    term_start = date.fromisoformat(args.term_start)
    if args.term_end:
        term_end = date.fromisoformat(args.term_end)
    else:
        term_end = date(term_start.year + 1, term_start.month, term_start.day) - timedelta(days=1)
    renewal_date = term_end + timedelta(days=1)
    slug = re.sub(r"[^a-z0-9]+", "-", args.account.lower()).strip("-")
    return DocConfig(
        account=args.account,
        supplier=args.supplier,
        product=args.product,
        base_price=args.base_price,
        quantity=args.quantity,
        currency=args.currency,
        uplift=args.uplift,
        amendment_uplift=args.amendment_uplift,
        term_start=term_start,
        term_end=term_end,
        renewal_date=renewal_date,
        notice_days=args.notice_days,
        slug=slug,
    )


# ─── Formatting helpers ───────────────────────────────────────────────────────

def fmt_money(amount: float, currency: str = "USD") -> str:
    symbols = {"USD": "$", "EUR": "€", "GBP": "£"}
    sym = symbols.get(currency, currency + " ")
    return f"{sym}{amount:,.2f}"


def fmt_date(d: date) -> str:
    return d.strftime("%B %d, %Y")


def fmt_date_iso(d: date) -> str:
    return d.isoformat()


def add_page_break(doc):
    doc.add_page_break()


def styled_heading(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h


def body(doc, text: str, bold_prefix: str = ""):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0)
    return p


def clause(doc, number: str, title: str, text: str):
    p = doc.add_paragraph()
    r = p.add_run(f"Section {number}. {title}. ")
    r.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.2)
    return p


def recital(doc, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(6)
    return p


def signature_block(doc, company: str, name: str, title: str, sign_date: date):
    doc.add_paragraph("")
    body(doc, company.upper())
    body(doc, f"By: /s/ {name}")
    body(doc, f"Name: {name}")
    body(doc, f"Title: {title}")
    body(doc, f"Date: {fmt_date(sign_date)}")


def info_table(doc, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph("")


# ─── Document 1: Master Subscription Agreement ────────────────────────────────

def create_msa(c: DocConfig, out_dir: Path) -> Path:
    doc = Document()
    monthly = c.base_price * c.quantity
    annual = monthly * 12

    styled_heading(doc, "MASTER SUBSCRIPTION AGREEMENT")

    body(doc, (
        f"This Master Subscription Agreement (\"Agreement\") is entered into as of "
        f"{fmt_date(c.term_start)}, by and between {c.supplier} (\"Supplier\") "
        f"and {c.account} (\"Customer\")."
    ))

    body(doc, (
        f"WHEREAS, Customer desires to subscribe to Supplier's {c.product} "
        f"(the \"Service\") for use by Customer's authorized users in connection with "
        f"Customer's internal business operations; and"
    ))
    body(doc, (
        "WHEREAS, Supplier desires to provide such subscription subject to the terms "
        "and conditions set forth below;"
    ))
    body(doc, "NOW, THEREFORE, in consideration of the mutual covenants herein, the parties agree as follows:")

    clause(doc, "1", "Definitions",
        "\"Authorized Users\" means Customer's employees and contractors permitted "
        "to access the Service. \"Order Form\" means a mutually executed document "
        "specifying edition, quantities, and pricing. \"Subscription Term\" means "
        "the Initial Term and any Renewal Terms. \"Confidential Information\" means "
        "any non-public information disclosed by one party to the other that is "
        "designated as confidential or that reasonably should be understood to be "
        "confidential given the nature of the information and circumstances of disclosure."
    )

    clause(doc, "2", "Subscription Grant",
        f"Subject to Customer's payment of all fees, Supplier grants Customer a "
        f"limited, non-exclusive, non-transferable right to access and use the "
        f"Service for Customer's internal business purposes during the Subscription "
        f"Term. Customer may not sublicense, resell, or permit use of the Service "
        f"by any third party without Supplier's prior written consent."
    )

    clause(doc, "3", "Subscription Term",
        f"The Initial Term commences on {fmt_date(c.term_start)} and continues "
        f"through {fmt_date(c.term_end)}. Thereafter, this Agreement will "
        f"automatically renew for successive twelve (12) month Renewal Terms unless "
        f"either party provides written notice of non-renewal at least {c.notice_days} "
        f"days prior to the end of the then-current term."
    )

    clause(doc, "4", "Fees and Payment",
        f"Customer will pay a monthly subscription fee of "
        f"{fmt_money(c.base_price, c.currency)} per authorized seat. As of the "
        f"Effective Date, Customer has subscribed for {c.quantity:,} seats, for a "
        f"total monthly commitment of {fmt_money(monthly, c.currency)} "
        f"({fmt_money(annual, c.currency)} annually). Undisputed invoices are due "
        f"net thirty (30) days from invoice date. Late payments accrue interest at "
        f"1.5% per month or the maximum rate permitted by law, whichever is lower."
    )

    clause(doc, "5", "Annual Price Adjustment",
        f"Beginning with the first Renewal Term and on each subsequent anniversary "
        f"of the Effective Date, the recurring subscription fees shall increase by "
        f"{c.uplift:.1f}% (the \"Annual Uplift\"). Supplier shall provide Customer "
        f"with at least {c.notice_days} days' prior written notice before the "
        f"applicable anniversary date. If Supplier fails to provide timely notice, "
        f"the Annual Uplift right for that period shall be deferred to the following "
        f"anniversary."
    )

    clause(doc, "6", "Service Levels",
        "Supplier will use commercially reasonable efforts to maintain Service "
        "availability of at least 99.5% uptime measured monthly, excluding scheduled "
        "maintenance windows communicated at least 48 hours in advance. In the event "
        "of a material service outage exceeding four (4) consecutive hours, Customer "
        "may request a service credit equal to five percent (5%) of the monthly "
        "subscription fee for the affected month."
    )

    clause(doc, "7", "Data Protection",
        "Supplier will maintain reasonable administrative, technical, and physical "
        "safeguards to protect Customer Data. Supplier will comply with applicable "
        "data protection laws and will notify Customer within 72 hours of discovering "
        "a security incident that materially affects Customer Data."
    )

    clause(doc, "8", "Confidentiality",
        "Each party agrees not to disclose the other party's Confidential Information "
        "to any third party without prior written consent, and to use Confidential "
        "Information only as necessary to perform obligations under this Agreement. "
        "Each party will protect Confidential Information with at least the same "
        "degree of care it uses for its own similar information, but no less than "
        "reasonable care. Obligations survive termination for three (3) years."
    )

    clause(doc, "9", "Intellectual Property",
        "Supplier retains all right, title, and interest in and to the Service, "
        "including all underlying software, algorithms, models, and documentation. "
        "Customer retains all right, title, and interest in and to Customer Data. "
        "Nothing in this Agreement transfers ownership of either party's IP to "
        "the other."
    )

    clause(doc, "10", "Limitation of Liability",
        "EXCEPT FOR BREACHES OF SECTION 7 (DATA PROTECTION) OR SECTION 8 "
        "(CONFIDENTIALITY), OR FOR INDEMNIFICATION OBLIGATIONS, NEITHER PARTY'S "
        "AGGREGATE LIABILITY UNDER THIS AGREEMENT SHALL EXCEED THE TOTAL FEES "
        "PAID OR PAYABLE BY CUSTOMER DURING THE TWELVE (12) MONTH PERIOD "
        "IMMEDIATELY PRECEDING THE CLAIM. IN NO EVENT SHALL EITHER PARTY BE "
        "LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, OR CONSEQUENTIAL DAMAGES."
    )

    clause(doc, "11", "Termination",
        "Either party may terminate this Agreement upon written notice if the other "
        "party materially breaches this Agreement and fails to cure such breach "
        "within thirty (30) days of receiving written notice. Upon termination, "
        "Customer's right to access the Service ceases and all unpaid fees become "
        "immediately due."
    )

    clause(doc, "12", "Governing Law and Disputes",
        "This Agreement will be governed by the laws of the State of Delaware, "
        "without regard to conflict-of-law principles. The parties agree to first "
        "attempt good-faith negotiation to resolve any dispute. If unresolved "
        "within 30 days, disputes shall be settled by binding arbitration under "
        "JAMS rules in San Francisco, California."
    )

    clause(doc, "13", "Amendment Priority",
        "Order Forms, commercial amendments, and renewal schedules executed after "
        "the Effective Date may modify the commercial terms herein. In the event "
        "of a conflict between this Agreement and an Order Form or Amendment, the "
        "later-dated instrument shall control solely with respect to the conflicting "
        "commercial terms."
    )

    clause(doc, "14", "Entire Agreement",
        "This Agreement, together with all Order Forms, amendments, and exhibits, "
        "constitutes the entire agreement between the parties with respect to the "
        "subject matter and supersedes all prior or contemporaneous negotiations "
        "and communications, whether written or oral."
    )

    add_page_break(doc)
    body(doc, "IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.")
    body(doc, "")
    signature_block(doc, c.supplier, "Sarah Mitchell", "VP, Commercial Operations",
                    c.term_start - timedelta(days=3))
    body(doc, "")
    signature_block(doc, c.account, "David Chen", "Chief Financial Officer",
                    c.term_start - timedelta(days=2))

    path = out_dir / f"{c.slug}-master-subscription-agreement-v1.docx"
    doc.save(path)
    print(f"  Created: {path.name}")
    return path


# ─── Document 2: Order Form ───────────────────────────────────────────────────

def create_order_form(c: DocConfig, out_dir: Path) -> Path:
    doc = Document()
    monthly = c.base_price * c.quantity
    annual = monthly * 12

    styled_heading(doc, "ORDER FORM")
    styled_heading(doc, f"Order Form No. OF-{c.term_start.year}-001", level=2)

    body(doc, (
        f"This Order Form (\"Order Form\") is entered into as of {fmt_date(c.term_start)} "
        f"between {c.supplier} (\"Supplier\") and {c.account} (\"Customer\") and "
        f"is governed by the Master Subscription Agreement between the parties dated "
        f"{fmt_date(c.term_start)} (the \"Agreement\")."
    ))

    styled_heading(doc, "1. Service Details", level=2)
    info_table(doc, [
        ("Product", c.product),
        ("Edition", "Enterprise"),
        ("Subscription Term", f"{fmt_date_iso(c.term_start)} to {fmt_date_iso(c.term_end)}"),
        ("Number of Seats", f"{c.quantity:,}"),
        ("Price per Seat (monthly)", fmt_money(c.base_price, c.currency)),
        ("Monthly Commitment", fmt_money(monthly, c.currency)),
        ("Annual Commitment", fmt_money(annual, c.currency)),
        ("Billing Frequency", "Monthly, in arrears"),
        ("Currency", c.currency),
        ("Payment Terms", "Net 30 days from invoice date"),
    ])

    styled_heading(doc, "2. Pricing Schedule", level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for cell, text in zip(hdr.cells, ["Period", "Seats", "Price/Seat/Mo", "Monthly Total"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
    year1_price = c.base_price
    for i in range(1, 4):
        row = table.rows[i]
        period_start = date(c.term_start.year + (i - 1), c.term_start.month, c.term_start.day)
        period_end = date(c.term_start.year + i, c.term_start.month, c.term_start.day) - timedelta(days=1)
        price = year1_price * ((1 + c.uplift / 100) ** (i - 1))
        row.cells[0].text = f"Year {i} ({fmt_date_iso(period_start)} – {fmt_date_iso(period_end)})"
        row.cells[1].text = f"{c.quantity:,}"
        row.cells[2].text = fmt_money(price, c.currency)
        row.cells[3].text = fmt_money(price * c.quantity, c.currency)
    note_row = table.rows[4]
    note_row.cells[0].text = "Note"
    note_row.cells[1].merge(note_row.cells[2]).merge(note_row.cells[3])
    note_row.cells[1].text = (
        f"Year 2 and Year 3 pricing reflects the {c.uplift:.1f}% Annual Uplift "
        f"per Section 5 of the Agreement."
    )
    doc.add_paragraph("")

    styled_heading(doc, "3. Annual Price Adjustment Confirmation", level=2)
    body(doc, (
        f"The parties confirm that, pursuant to Section 5 of the Agreement, the "
        f"recurring subscription fees shall increase by {c.uplift:.1f}% on each "
        f"anniversary of the Subscription Term Start Date. Supplier will provide "
        f"Customer with written notice at least {c.notice_days} days before each "
        f"anniversary."
    ))

    styled_heading(doc, "4. Implementation and Onboarding", level=2)
    body(doc, (
        "Supplier will provide standard onboarding services at no additional charge, "
        "including up to eight (8) hours of implementation support, configuration "
        "assistance, and user training sessions. Onboarding will be scheduled "
        "within fifteen (15) business days of the Effective Date."
    ))

    styled_heading(doc, "5. Special Terms", level=2)
    body(doc, (
        f"Customer is entitled to designate up to ten percent (10%) of licensed "
        f"seats as read-only viewer seats at no additional charge. Any seats in "
        f"excess of {c.quantity:,} will be charged at the then-current per-seat rate "
        f"on a pro-rated basis."
    ))

    add_page_break(doc)
    body(doc, "This Order Form is hereby agreed and accepted by the authorized representatives of each party.")
    body(doc, "")
    signature_block(doc, c.supplier, "Sarah Mitchell", "VP, Commercial Operations",
                    c.term_start - timedelta(days=2))
    body(doc, "")
    signature_block(doc, c.account, "David Chen", "Chief Financial Officer",
                    c.term_start - timedelta(days=1))

    path = out_dir / f"{c.slug}-order-form-v1.docx"
    doc.save(path)
    print(f"  Created: {path.name}")
    return path


# ─── Document 3: Commercial Amendment ─────────────────────────────────────────

def create_amendment(c: DocConfig, out_dir: Path, version: int = 1) -> Path:
    doc = Document()
    amend_date = c.term_start + timedelta(days=180)   # ~6 months into term
    eff_date = amend_date + timedelta(days=14)
    new_monthly = c.base_price * c.quantity * (1 + c.amendment_uplift / 100)

    styled_heading(doc, f"COMMERCIAL AMENDMENT NO. {version}")
    styled_heading(doc, f"To the Master Subscription Agreement dated {fmt_date(c.term_start)}", level=2)

    body(doc, (
        f"This Commercial Amendment No. {version} (\"Amendment\") is entered into "
        f"as of {fmt_date(amend_date)}, by and between {c.supplier} "
        f"(\"Supplier\") and {c.account} (\"Customer\"), and amends the Master "
        f"Subscription Agreement dated {fmt_date(c.term_start)} and Order Form "
        f"No. OF-{c.term_start.year}-001 (collectively, the \"Agreement\")."
    ))

    styled_heading(doc, "Recitals", level=2)
    recital(doc, (
        "WHEREAS, the parties desire to update certain commercial terms of the "
        "Agreement to reflect current market conditions, platform expansion, and "
        "Customer's increased usage requirements;"
    ))
    recital(doc, (
        f"WHEREAS, the parties have agreed that a revised annual price adjustment "
        f"rate of {c.amendment_uplift:.1f}% better reflects the expanded scope of "
        f"services being provided under the Agreement;"
    ))
    recital(doc, "NOW, THEREFORE, in consideration of the mutual covenants set forth herein, the parties agree as follows:")

    styled_heading(doc, "Amendment Terms", level=2)

    clause(doc, "A.1", "Supersession of Prior Pricing Uplift Clause",
        f"Section 5 (Annual Price Adjustment) of the Master Subscription Agreement "
        f"and any corresponding pricing language in Order Form No. OF-{c.term_start.year}-001 "
        f"are hereby deleted in their entirety and replaced with the following: "
        f"\"Beginning on {fmt_date(eff_date)} and on each subsequent annual "
        f"anniversary date, the recurring subscription fees shall increase by "
        f"{c.amendment_uplift:.1f}% (the 'Revised Annual Uplift'). Supplier shall "
        f"provide Customer with at least {c.notice_days} days' prior written "
        f"notice before each anniversary.\""
    )

    clause(doc, "A.2", "Effective Date of Revised Uplift",
        f"The Revised Annual Uplift of {c.amendment_uplift:.1f}% shall take effect "
        f"beginning {fmt_date(eff_date)}. All invoices issued on or after "
        f"{fmt_date(eff_date)} shall reflect the Revised Annual Uplift rate for "
        f"the applicable renewal period. The prior Annual Uplift rate of "
        f"{c.uplift:.1f}% set forth in the Agreement shall no longer apply."
    )

    clause(doc, "A.3", "Updated Pricing Schedule",
        f"For the avoidance of doubt, effective {fmt_date(eff_date)}, the monthly "
        f"subscription fee for {c.quantity:,} seats shall be calculated at "
        f"{fmt_money(c.base_price * (1 + c.amendment_uplift / 100), c.currency)} "
        f"per seat, for a total monthly commitment of "
        f"{fmt_money(new_monthly, c.currency)}."
    )

    clause(doc, "A.4", "Expanded Service Scope",
        f"Customer's subscription is expanded to include access to all standard "
        f"modules released by Supplier during the Subscription Term at no additional "
        f"charge, subject to Customer remaining current on all payment obligations."
    )

    clause(doc, "A.5", "Controlling Terms",
        "In the event of any conflict between this Amendment and any prior version "
        "of the Agreement (including the original Master Subscription Agreement or "
        "any prior amendments), the terms of this Amendment shall control. All "
        "other terms and conditions of the Agreement not expressly modified herein "
        "remain in full force and effect."
    )

    clause(doc, "A.6", "Ratification",
        "Each party hereby ratifies and confirms the Agreement, as amended by "
        "this Amendment, and acknowledges that the Agreement, as so amended, "
        "is in full force and effect."
    )

    add_page_break(doc)
    body(doc, "IN WITNESS WHEREOF, the parties have executed this Amendment as of the date first written above.")
    body(doc, "")
    signature_block(doc, c.supplier, "Sarah Mitchell", "VP, Commercial Operations", amend_date)
    body(doc, "")
    signature_block(doc, c.account, "David Chen", "Chief Financial Officer",
                    amend_date + timedelta(days=1))

    path = out_dir / f"{c.slug}-commercial-amendment-v{version}.docx"
    doc.save(path)
    print(f"  Created: {path.name}")
    return path


# ─── Document 4: Renewal Notice ────────────────────────────────────────────────

def create_renewal_notice(c: DocConfig, out_dir: Path, version: int = 1) -> Path:
    doc = Document()
    notice_date = c.term_end - timedelta(days=c.notice_days + 5)  # sent 5 days before deadline
    new_price_per_seat = c.base_price * (1 + c.amendment_uplift / 100)
    new_monthly = new_price_per_seat * c.quantity
    new_annual = new_monthly * 12

    styled_heading(doc, "RENEWAL NOTICE AND PRICING CONFIRMATION")

    body(doc, f"Date: {fmt_date(notice_date)}")
    body(doc, "")
    body(doc, f"To: {c.account}")
    body(doc, "Attn: Chief Financial Officer")
    body(doc, "")
    body(doc, f"From: {c.supplier} — Commercial Operations")
    body(doc, f"Re: Subscription Renewal — {c.product}")
    body(doc, "")

    styled_heading(doc, "1. Notice of Renewal", level=2)
    body(doc, (
        f"Pursuant to Section 3 (Subscription Term) of the Master Subscription "
        f"Agreement dated {fmt_date(c.term_start)} (the \"Agreement\"), "
        f"{c.supplier} hereby provides formal notice that the subscription for "
        f"{c.product} will automatically renew for a twelve (12) month Renewal "
        f"Term commencing on {fmt_date(c.renewal_date)}."
    ))

    styled_heading(doc, "2. Revised Pricing Schedule", level=2)
    body(doc, (
        f"As set forth in Commercial Amendment No. 1 dated "
        f"{fmt_date(c.term_start + timedelta(days=180))}, the Annual Uplift "
        f"rate of {c.amendment_uplift:.1f}% applies to the Renewal Term. "
        f"The following pricing will take effect on {fmt_date(c.renewal_date)}:"
    ))

    info_table(doc, [
        ("Renewal Term", f"{fmt_date_iso(c.renewal_date)} to "
                         f"{date(c.renewal_date.year + 1, c.renewal_date.month, c.renewal_date.day) - timedelta(days=1)}"),
        ("Product", c.product),
        ("Number of Seats", f"{c.quantity:,}"),
        ("Price per Seat (monthly)", fmt_money(new_price_per_seat, c.currency)),
        ("Total Monthly Commitment", fmt_money(new_monthly, c.currency)),
        ("Total Annual Commitment", fmt_money(new_annual, c.currency)),
        ("Annual Uplift Rate Applied", f"{c.amendment_uplift:.1f}%"),
        ("Governing Amendment", "Commercial Amendment No. 1"),
    ])

    styled_heading(doc, "3. Basis for Pricing", level=2)
    body(doc, (
        f"The revised pricing reflects the {c.amendment_uplift:.1f}% Annual Uplift "
        f"rate agreed in Commercial Amendment No. 1. The original Annual Uplift "
        f"rate of {c.uplift:.1f}% set forth in the Master Subscription Agreement "
        f"was superseded by Commercial Amendment No. 1 and does not apply to "
        f"this Renewal Term."
    ))

    styled_heading(doc, "4. Non-Renewal", level=2)
    body(doc, (
        f"If Customer does not wish to renew, Customer must provide written notice "
        f"of non-renewal to {c.supplier} no later than {fmt_date(c.term_end - timedelta(days=c.notice_days))}. "
        f"If no such notice is received by that date, the subscription will "
        f"automatically renew on the terms described herein."
    ))

    styled_heading(doc, "5. Next Steps", level=2)
    body(doc, (
        "No action is required if Customer wishes to renew under the pricing and "
        "terms described in this notice. Invoices for the Renewal Term will be "
        "issued automatically at the updated pricing. If you have questions or "
        f"require an updated Order Form, please contact your {c.supplier} account "
        f"manager."
    ))

    body(doc, "")
    body(doc, "This notice is provided in accordance with the Agreement.")
    body(doc, "")
    body(doc, f"Sincerely,")
    body(doc, "")
    body(doc, "Sarah Mitchell")
    body(doc, "VP, Commercial Operations")
    body(doc, c.supplier)
    body(doc, f"Date: {fmt_date(notice_date)}")

    path = out_dir / f"{c.slug}-renewal-notice-v{version}.docx"
    doc.save(path)
    print(f"  Created: {path.name}")
    return path


# ─── CLI ─────────────────────────────────────────────────────────────────────

GENERATORS = {
    "msa": create_msa,
    "order_form": create_order_form,
    "amendment": create_amendment,
    "renewal_notice": create_renewal_notice,
}


def  parse_args():
    p = argparse.ArgumentParser(
        description="Generate demo contract DOCX files for any account.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    p.add_argument("--account", required=True, help="Customer company name")
    p.add_argument("--product", required=True, help="Product/subscription name")
    p.add_argument("--base-price", type=float, required=True,
                   help="Per-unit monthly price (e.g. 8000)")
    p.add_argument("--quantity", type=int, default=100,
                   help="Number of seats/units (default: 100)")
    p.add_argument("--currency", default="USD", choices=["USD", "EUR", "GBP"],
                   help="Currency (default: USD)")
    p.add_argument("--uplift", type=float, default=5.0,
                   help="MSA annual uplift %% (default: 5.0)")
    p.add_argument("--amendment-uplift", type=float, default=None,
                   help="Amendment override uplift %% (default: uplift + 3)")
    p.add_argument("--term-start", default=date.today().isoformat(),
                   help="Term start date YYYY-MM-DD (default: today)")
    p.add_argument("--term-end", default=None,
                   help="Term end date YYYY-MM-DD (default: term-start + 1 year - 1 day)")
    p.add_argument("--notice-days", type=int, default=30,
                   help="Renewal notice window in days (default: 30)")
    p.add_argument("--supplier", default="Conga Software, Inc.",
                   help="Supplier name (default: Conga Software, Inc.)")
    p.add_argument("--types", nargs="+",
                   choices=list(GENERATORS.keys()) + ["all"],
                   default=["all"],
                   help="Document types to generate (default: all)")
    p.add_argument("--output-dir", default=None,
                   help="Output directory (default: demo/<account-slug>/)")
    return p.parse_args()


def main():
    args = parse_args()

    if args.amendment_uplift is None:
        args.amendment_uplift = args.uplift + 3.0

    cfg = make_config(args)

    out_dir = Path(args.output_dir) if args.output_dir else (
        Path(__file__).resolve().parent.parent / "demo" / cfg.slug
    )
    out_dir.mkdir(parents=True, exist_ok=True)

    types = list(GENERATORS.keys()) if "all" in args.types else args.types

    print(f"\nGenerating documents for: {cfg.account}")
    print(f"  Product:          {cfg.product}")
    print(f"  Base price:       {fmt_money(cfg.base_price, cfg.currency)}/seat/mo × {cfg.quantity:,} seats")
    print(f"  MSA uplift:       {cfg.uplift:.1f}%")
    print(f"  Amendment uplift: {cfg.amendment_uplift:.1f}%  (overrides MSA)")
    print(f"  Term:             {cfg.term_start} → {cfg.term_end}")
    print(f"  Output:           {out_dir}")
    print()

    generated = []
    for doc_type in types:
        fn = GENERATORS[doc_type]
        if doc_type == "amendment":
            path = fn(cfg, out_dir, version=1)
        elif doc_type == "renewal_notice":
            path = fn(cfg, out_dir, version=1)
        else:
            path = fn(cfg, out_dir)
        generated.append(path)

    print(f"\nDone. {len(generated)} file(s) written to {out_dir}")
    print("\nUpload order for the leakage demo:")
    order = [
        ("msa",            "subscription-agreement", "master agreement"),
        ("order_form",     "order-form",             "order form"),
        ("amendment",      "amendment",              "amendment"),
        ("renewal_notice", "renewal-notice",         "renewal notice"),
    ]
    n = 0
    for doc_type, name_fragment, label in order:
        if doc_type not in types:
            continue
        match = [p for p in generated if name_fragment in p.name]
        if match:
            n += 1
            print(f"  {n}. {match[0].name}  →  document type: {label}")


if __name__ == "__main__":
    main()
