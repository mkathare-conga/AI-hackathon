from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUTPUT_DIR = Path(__file__).resolve().parent.parent / "demo" / "dummy_docs"
SIMILAR_COMPANY_OUTPUT_DIR = Path(__file__).resolve().parent.parent / "demo" / "dummy_docs_similar_companies"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(8)


def add_clause(doc: Document, number: int, title: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"{number}. {title}. ")
    run.bold = True
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(10)


def add_signature_block(doc: Document, supplier_signer: str, customer_name: str, customer_signer: str, signed_date: str) -> None:
    add_body(doc, "")
    add_body(doc, "CONGA SOFTWARE, INC.")
    add_body(doc, f"By: /s/ {supplier_signer}")
    add_body(doc, f"Name: {supplier_signer}")
    add_body(doc, "Title: VP, Commercial Operations")
    add_body(doc, f"Date: {signed_date}")

    add_body(doc, "")
    add_body(doc, customer_name.upper())
    add_body(doc, f"By: /s/ {customer_signer}")
    add_body(doc, f"Name: {customer_signer}")
    add_body(doc, "Title: Chief Financial Officer")
    add_body(doc, f"Date: {signed_date}")


def build_msa(doc_info: dict[str, str]) -> Document:
    doc = Document()
    title = doc.add_heading(doc_info["title"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body(
        doc,
        (
            f"This Master Subscription Agreement is entered into as of {doc_info['effective_date']} by and between "
            f"Conga Software, Inc. and {doc_info['customer']} for the use of {doc_info['product']}."
        ),
    )
    add_clause(
        doc,
        1,
        "Subscription Term",
        (
            f"The initial term begins on {doc_info['term_start']} and continues through {doc_info['term_end']}. "
            "The subscription renews automatically for successive twelve-month periods unless either party provides timely written notice of non-renewal."
        ),
    )
    add_clause(
        doc,
        2,
        "Fees and Invoicing",
        (
            f"Customer will be invoiced monthly for {doc_info['quantity']} units at {doc_info['currency']} {doc_info['unit_price']} per unit. "
            "Undisputed invoices are due net thirty (30) days from receipt."
        ),
    )
    add_clause(
        doc,
        3,
        "Renewal Pricing Adjustment",
        (
            f"Beginning with the first renewal term, recurring subscription charges are subject to a {doc_info['uplift']} annual uplift. "
            f"Supplier may implement the {doc_info['uplift']} annual uplift by providing Customer at least {doc_info['notice_days']} days notice before the applicable renewal anniversary."
        ),
    )
    add_clause(
        doc,
        4,
        "Precedence",
        "Later order forms, amendments, and renewal notices may supersede this agreement solely with respect to pricing, quantities, and renewal mechanics if they expressly state that they control.",
    )
    add_signature_block(doc, doc_info["supplier_signer"], doc_info["customer"], doc_info["customer_signer"], doc_info["signed_date"])
    return doc


def build_order_form(doc_info: dict[str, str]) -> Document:
    doc = Document()
    title = doc.add_heading(doc_info["title"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body(doc, f"Customer: {doc_info['customer']}")
    add_body(doc, f"Product: {doc_info['product']}")
    add_body(doc, f"Commercial Start Date: {doc_info['term_start']}")
    add_body(doc, f"Committed End Date: {doc_info['term_end']}")
    add_clause(
        doc,
        1,
        "Subscription Commitment",
        f"Customer is purchasing {doc_info['quantity']} subscribed seats or units at a base unit price of {doc_info['currency']} {doc_info['unit_price']}.",
    )
    add_clause(
        doc,
        2,
        "Commercial Override",
        (
            f"For commercial pricing only, this order form supersedes prior pricing language and sets the next renewal uplift at {doc_info['uplift']}. "
            f"Supplier may implement the {doc_info['uplift']} uplift by providing Customer at least {doc_info['notice_days']} days notice before renewal."
        ),
    )
    add_clause(
        doc,
        3,
        "Billing",
        "Supplier will invoice monthly in arrears based on the committed subscription quantities unless otherwise stated in a later amendment.",
    )
    add_signature_block(doc, doc_info["supplier_signer"], doc_info["customer"], doc_info["customer_signer"], doc_info["signed_date"])
    return doc


def build_amendment(doc_info: dict[str, str]) -> Document:
    doc = Document()
    title = doc.add_heading(doc_info["title"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, f"to the agreement effective {doc_info['prior_agreement_date']}", level=2)

    add_body(
        doc,
        (
            f"This amendment is entered into as of {doc_info['effective_date']} by and between Conga Software, Inc. and {doc_info['customer']}. "
            "It amends the commercial pricing terms solely as described below."
        ),
    )
    add_clause(
        doc,
        1,
        "Amended Renewal Pricing",
        (
            f"Beginning with the renewal term commencing on {doc_info['renewal_effective_date']}, recurring subscription fees are subject to a {doc_info['uplift']} annual uplift. "
            f"Supplier may implement the {doc_info['uplift']} annual uplift by providing Customer at least {doc_info['notice_days']} days written notice before the applicable renewal anniversary."
        ),
    )
    add_clause(
        doc,
        2,
        "Supersession",
        (
            f"This amendment supersedes and replaces any prior {doc_info['prior_uplift']} uplift references for the covered product line, "
            "but all other terms of the underlying agreement remain unchanged and in full force."
        ),
    )
    add_clause(
        doc,
        3,
        "Commercial Rationale",
        doc_info["rationale"],
    )
    add_signature_block(doc, doc_info["supplier_signer"], doc_info["customer"], doc_info["customer_signer"], doc_info["signed_date"])
    return doc


def build_renewal_notice(doc_info: dict[str, str]) -> Document:
    doc = Document()
    title = doc.add_heading(doc_info["title"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    currency = doc_info.get("currency", "USD")

    add_body(doc, f"Date: {doc_info['effective_date']}")
    add_body(doc, f"To: {doc_info['customer']}")
    add_body(doc, f"Re: Renewal pricing notice for {doc_info['product']}")
    add_body(
        doc,
        (
            f"This letter serves as formal notice that the annual renewal pricing adjustment will take effect on {doc_info['renewal_effective_date']}. "
            f"The current governing commercial position is a {doc_info['uplift']} annual uplift with a {doc_info['notice_days']}-day notice requirement."
        ),
    )
    add_clause(
        doc,
        1,
        "Updated Pricing",
        (
            f"Current monthly fee: {currency} {doc_info['current_monthly_fee']}. "
            f"Adjusted monthly fee after applying the {doc_info['uplift']} uplift: {currency} {doc_info['new_monthly_fee']}."
        ),
    )
    add_clause(
        doc,
        2,
        "Notice Timing",
        (
            f"This notice is intended to satisfy the {doc_info['notice_days']}-day advance notice requirement before the renewal date of {doc_info['renewal_effective_date']}. "
            f"The internal deadline to send the notice was {doc_info['deadline_date']}."
        ),
    )
    add_clause(doc, 3, "Operational Note", doc_info["rationale"])
    add_body(doc, "Sincerely,")
    add_body(doc, doc_info["supplier_signer"])
    add_body(doc, "VP, Commercial Operations")
    return doc


def build_nda(doc_info: dict[str, str]) -> Document:
    doc = Document()
    title = doc.add_heading(doc_info["title"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body(
        doc,
        (
            f"This Mutual Non-Disclosure Agreement is entered into as of {doc_info['effective_date']} between Conga Software, Inc. and {doc_info['customer']} "
            "for the purpose of evaluating a potential commercial relationship."
        ),
    )
    add_clause(doc, 1, "Confidential Information", "Each party will protect the other party's confidential information using reasonable care and will not disclose it to third parties except as permitted by this agreement.")
    add_clause(doc, 2, "Permitted Use", "Confidential information may be used solely for evaluating and supporting the parties' business relationship and may not be used for any competitive purpose.")
    add_clause(doc, 3, "No Commercial Override", "This NDA does not establish pricing, renewal, uplift, billing, or other commercial obligations and does not modify any subscription agreement between the parties.")
    add_signature_block(doc, doc_info["supplier_signer"], doc_info["customer"], doc_info["customer_signer"], doc_info["signed_date"])
    return doc


def build_internal_memo(doc_info: dict[str, str]) -> Document:
    doc = Document()
    title = doc.add_heading(doc_info["title"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body(doc, f"Prepared for: {doc_info['customer']}")
    add_body(doc, f"Document date: {doc_info['effective_date']}")
    add_body(
        doc,
        (
            f"This internal planning memo relates to {doc_info['product']} operations, adoption planning, and customer-success execution. "
            "It is not a signed commercial instrument and does not modify contractual pricing or renewal rights."
        ),
    )
    add_clause(doc, 1, "Operational Context", doc_info["context"])
    add_clause(doc, 2, "Non-Controlling Status", "This memo is for internal planning only. It does not establish pricing, uplift, notice obligations, billing mechanics, or document precedence.")
    add_clause(doc, 3, "Next Steps", doc_info["rationale"])
    add_body(doc, f"Author: {doc_info['supplier_signer']}")
    return doc


DOCUMENTS = [
    {
        "builder": build_msa,
        "file_name": "aldera-master-subscription-agreement-v1.docx",
        "title": "MASTER SUBSCRIPTION AGREEMENT",
        "customer": "Aldera Manufacturing Group",
        "product": "Conga CLM Enterprise",
        "effective_date": "January 10, 2025",
        "term_start": "January 10, 2025",
        "term_end": "January 9, 2026",
        "quantity": "850",
        "currency": "USD",
        "unit_price": "125.00",
        "uplift": "4%",
        "notice_days": "30",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Monica Reeves",
        "signed_date": "January 10, 2025",
    },
    {
        "builder": build_order_form,
        "file_name": "aldera-enterprise-order-form-v2.docx",
        "title": "ENTERPRISE ORDER FORM",
        "customer": "Aldera Manufacturing Group",
        "product": "Conga CLM Enterprise",
        "term_start": "January 10, 2025",
        "term_end": "January 9, 2026",
        "quantity": "850",
        "currency": "USD",
        "unit_price": "129.00",
        "uplift": "5%",
        "notice_days": "30",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Monica Reeves",
        "signed_date": "March 5, 2025",
    },
    {
        "builder": build_amendment,
        "file_name": "aldera-commercial-amendment-v1.docx",
        "title": "COMMERCIAL AMENDMENT NO. 1",
        "customer": "Aldera Manufacturing Group",
        "effective_date": "November 20, 2025",
        "prior_agreement_date": "January 10, 2025",
        "renewal_effective_date": "January 10, 2026",
        "uplift": "7%",
        "prior_uplift": "5%",
        "notice_days": "45",
        "rationale": "The parties expanded the analytics footprint and agreed to revised renewal pricing to reflect the broader deployment and higher support tier.",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Monica Reeves",
        "signed_date": "November 20, 2025",
    },
    {
        "builder": build_renewal_notice,
        "file_name": "aldera-renewal-notice-2026.docx",
        "title": "RENEWAL PRICING NOTICE",
        "customer": "Aldera Manufacturing Group",
        "product": "Conga CLM Enterprise",
        "effective_date": "November 24, 2025",
        "renewal_effective_date": "January 10, 2026",
        "uplift": "7%",
        "notice_days": "45",
        "current_monthly_fee": "106250.00",
        "new_monthly_fee": "113687.50",
        "deadline_date": "November 26, 2025",
        "rationale": "The account team approved the increase after completing the annual value review and confirming executive sponsorship.",
        "supplier_signer": "Sarah Mitchell",
    },
    {
        "builder": build_nda,
        "file_name": "brightline-mutual-nda-v1.docx",
        "title": "MUTUAL NON-DISCLOSURE AGREEMENT",
        "customer": "Brightline Health Partners",
        "effective_date": "February 14, 2025",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Devon Price",
        "signed_date": "February 14, 2025",
    },
    {
        "builder": build_msa,
        "file_name": "cedar-health-master-agreement-v1.docx",
        "title": "SOFTWARE SUBSCRIPTION AGREEMENT",
        "customer": "Cedar Health Systems",
        "product": "Conga Composer Premium",
        "effective_date": "April 1, 2025",
        "term_start": "April 1, 2025",
        "term_end": "March 31, 2026",
        "quantity": "220",
        "currency": "USD",
        "unit_price": "215.00",
        "uplift": "3%",
        "notice_days": "21",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Ethan Morales",
        "signed_date": "April 1, 2025",
    },
    {
        "builder": build_amendment,
        "file_name": "cedar-health-pricing-amendment-v2.docx",
        "title": "PRICING AMENDMENT NO. 2",
        "customer": "Cedar Health Systems",
        "effective_date": "January 15, 2026",
        "prior_agreement_date": "April 1, 2025",
        "renewal_effective_date": "April 1, 2026",
        "uplift": "6%",
        "prior_uplift": "3%",
        "notice_days": "60",
        "rationale": "Additional template automation modules and expanded managed services support were added during the renewal planning cycle.",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Ethan Morales",
        "signed_date": "January 15, 2026",
    },
    {
        "builder": build_order_form,
        "file_name": "harbor-retail-order-form-v1.docx",
        "title": "ORDER FORM",
        "customer": "Harbor Retail Holdings",
        "product": "Conga Orchestrate Standard",
        "term_start": "May 5, 2025",
        "term_end": "May 4, 2026",
        "quantity": "140",
        "currency": "USD",
        "unit_price": "145.00",
        "uplift": "2%",
        "notice_days": "30",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Priya Natarajan",
        "signed_date": "May 5, 2025",
    },
    {
        "builder": build_renewal_notice,
        "file_name": "harbor-retail-renewal-notice-draft-v1.docx",
        "title": "DRAFT RENEWAL NOTICE",
        "customer": "Harbor Retail Holdings",
        "product": "Conga Orchestrate Standard",
        "effective_date": "April 2, 2026",
        "renewal_effective_date": "May 5, 2026",
        "uplift": "2%",
        "notice_days": "30",
        "current_monthly_fee": "20300.00",
        "new_monthly_fee": "20706.00",
        "deadline_date": "April 5, 2026",
        "rationale": "This is a draft only and is awaiting sales leadership approval before the notice is sent to the customer.",
        "supplier_signer": "Sarah Mitchell",
    },
    {
        "builder": build_amendment,
        "file_name": "northstar-operations-commercial-amendment-v3.docx",
        "title": "COMMERCIAL AMENDMENT NO. 3",
        "customer": "Northstar Operations Ltd.",
        "effective_date": "August 18, 2025",
        "prior_agreement_date": "June 1, 2024",
        "renewal_effective_date": "September 1, 2025",
        "uplift": "9%",
        "prior_uplift": "6%",
        "notice_days": "75",
        "rationale": "The customer adopted the premium workflow governance package and the amendment supersedes earlier pricing references for the next renewal cycle.",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Jordan Kim",
        "signed_date": "August 18, 2025",
    },
]


SIMILAR_COMPANY_DOCUMENTS = [
    {
        "builder": build_msa,
        "file_name": "aldera-industrial-master-subscription-agreement-v1.docx",
        "title": "MASTER SUBSCRIPTION AGREEMENT",
        "customer": "Aldera Industrial Systems",
        "product": "Conga CLM Enterprise",
        "effective_date": "February 3, 2025",
        "term_start": "February 3, 2025",
        "term_end": "February 2, 2026",
        "quantity": "900",
        "currency": "USD",
        "unit_price": "121.00",
        "uplift": "4%",
        "notice_days": "30",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Lena Carter",
        "signed_date": "February 3, 2025",
    },
    {
        "builder": build_order_form,
        "file_name": "aldera-industrial-order-form-v2.docx",
        "title": "ENTERPRISE ORDER FORM",
        "customer": "Aldera Industrial Systems",
        "product": "Conga CLM Enterprise",
        "term_start": "February 3, 2025",
        "term_end": "February 2, 2026",
        "quantity": "900",
        "currency": "USD",
        "unit_price": "126.00",
        "uplift": "5%",
        "notice_days": "30",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Lena Carter",
        "signed_date": "April 10, 2025",
    },
    {
        "builder": build_amendment,
        "file_name": "aldera-industrial-commercial-amendment-v3.docx",
        "title": "COMMERCIAL AMENDMENT NO. 3",
        "customer": "Aldera Industrial Systems",
        "effective_date": "December 1, 2025",
        "prior_agreement_date": "February 3, 2025",
        "renewal_effective_date": "February 3, 2026",
        "uplift": "8%",
        "prior_uplift": "5%",
        "notice_days": "60",
        "rationale": "The parties expanded the deployment to additional manufacturing sites and agreed that the later amendment supersedes the earlier 4% and 5% pricing references.",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Lena Carter",
        "signed_date": "December 1, 2025",
    },
    {
        "builder": build_renewal_notice,
        "file_name": "aldera-industrial-renewal-notice-2026.docx",
        "title": "RENEWAL PRICING NOTICE",
        "customer": "Aldera Industrial Systems",
        "product": "Conga CLM Enterprise",
        "effective_date": "December 3, 2025",
        "renewal_effective_date": "February 3, 2026",
        "uplift": "8%",
        "notice_days": "60",
        "currency": "USD",
        "current_monthly_fee": "108900.00",
        "new_monthly_fee": "117612.00",
        "deadline_date": "December 5, 2025",
        "rationale": "This notice aligns to the latest amendment and is intended to preserve the renewal uplift for the 2026 cycle.",
        "supplier_signer": "Sarah Mitchell",
    },
    {
        "builder": build_internal_memo,
        "file_name": "aldera-industrial-qbr-memo-v1.docx",
        "title": "INTERNAL CUSTOMER SUCCESS MEMO",
        "customer": "Aldera Industrial Systems",
        "product": "Conga CLM Enterprise",
        "effective_date": "December 8, 2025",
        "context": "The quarterly business review highlighted strong adoption in two plants, an upcoming ERP cleanup, and interest in workflow automation training.",
        "rationale": "Coordinate with customer success and billing operations after renewal planning is finalized.",
        "supplier_signer": "Sarah Mitchell",
    },
    {
        "builder": build_msa,
        "file_name": "cedar-ridge-health-master-agreement-v1.docx",
        "title": "SOFTWARE SUBSCRIPTION AGREEMENT",
        "customer": "Cedar Ridge Health Group",
        "product": "Conga Composer Premium",
        "effective_date": "March 12, 2025",
        "term_start": "March 12, 2025",
        "term_end": "March 11, 2026",
        "quantity": "240",
        "currency": "USD",
        "unit_price": "208.00",
        "uplift": "3%",
        "notice_days": "21",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Marisol Vega",
        "signed_date": "March 12, 2025",
    },
    {
        "builder": build_amendment,
        "file_name": "cedar-ridge-health-pricing-amendment-v1.docx",
        "title": "PRICING AMENDMENT NO. 1",
        "customer": "Cedar Ridge Health Group",
        "effective_date": "January 6, 2026",
        "prior_agreement_date": "March 12, 2025",
        "renewal_effective_date": "March 12, 2026",
        "uplift": "4%",
        "prior_uplift": "3%",
        "notice_days": "30",
        "rationale": "The first pricing amendment updated support coverage and replaced the original uplift clause for the upcoming renewal.",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Marisol Vega",
        "signed_date": "January 6, 2026",
    },
    {
        "builder": build_order_form,
        "file_name": "cedar-ridge-health-renewal-schedule-v2.docx",
        "title": "RENEWAL SCHEDULE",
        "customer": "Cedar Ridge Health Group",
        "product": "Conga Composer Premium",
        "term_start": "March 12, 2026",
        "term_end": "March 11, 2027",
        "quantity": "240",
        "currency": "USD",
        "unit_price": "220.50",
        "uplift": "6%",
        "notice_days": "14",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Marisol Vega",
        "signed_date": "February 18, 2026",
    },
    {
        "builder": build_renewal_notice,
        "file_name": "cedar-ridge-health-renewal-notice-draft-v1.docx",
        "title": "DRAFT RENEWAL NOTICE",
        "customer": "Cedar Ridge Health Group",
        "product": "Conga Composer Premium",
        "effective_date": "February 24, 2026",
        "renewal_effective_date": "March 12, 2026",
        "uplift": "6%",
        "notice_days": "14",
        "currency": "USD",
        "current_monthly_fee": "49920.00",
        "new_monthly_fee": "52915.20",
        "deadline_date": "February 27, 2026",
        "rationale": "The notice is still in draft because finance is validating the final customer communication language.",
        "supplier_signer": "Sarah Mitchell",
    },
    {
        "builder": build_nda,
        "file_name": "cedar-ridge-health-mutual-nda-v1.docx",
        "title": "MUTUAL NON-DISCLOSURE AGREEMENT",
        "customer": "Cedar Ridge Health Group",
        "effective_date": "May 2, 2025",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Marisol Vega",
        "signed_date": "May 2, 2025",
    },
    {
        "builder": build_msa,
        "file_name": "harborline-retail-master-agreement-v1.docx",
        "title": "MASTER SUBSCRIPTION AGREEMENT",
        "customer": "Harborline Retail Group",
        "product": "Conga Orchestrate Standard",
        "effective_date": "June 9, 2025",
        "term_start": "June 9, 2025",
        "term_end": "June 8, 2026",
        "quantity": "160",
        "currency": "USD",
        "unit_price": "149.00",
        "uplift": "2%",
        "notice_days": "30",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Aisha Rahman",
        "signed_date": "June 9, 2025",
    },
    {
        "builder": build_amendment,
        "file_name": "harborline-retail-commercial-amendment-v2.docx",
        "title": "COMMERCIAL AMENDMENT NO. 2",
        "customer": "Harborline Retail Group",
        "effective_date": "April 15, 2026",
        "prior_agreement_date": "June 9, 2025",
        "renewal_effective_date": "June 9, 2026",
        "uplift": "5%",
        "prior_uplift": "2%",
        "notice_days": "45",
        "rationale": "The retailer added new regions and the amendment supersedes the original uplift language for the next renewal cycle.",
        "supplier_signer": "Sarah Mitchell",
        "customer_signer": "Aisha Rahman",
        "signed_date": "April 15, 2026",
    },
    {
        "builder": build_internal_memo,
        "file_name": "harborline-retail-operations-memo-v1.docx",
        "title": "RETAIL OPERATIONS MEMO",
        "customer": "Harborline Retail Group",
        "product": "Conga Orchestrate Standard",
        "effective_date": "April 22, 2026",
        "context": "Operations requested weekend enablement support, store manager training, and a rollout checklist for newly acquired store locations.",
        "rationale": "Track implementation readiness separately from renewal pricing and legal notice workflows.",
        "supplier_signer": "Sarah Mitchell",
    },
]


def write_documents(output_dir: Path, documents: list[dict[str, str]]) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)

    for doc_info in documents:
        builder = doc_info["builder"]
        output_path = output_dir / doc_info["file_name"]
        document = builder(doc_info)
        document.save(output_path)
        print(f"Created: {output_path.name}")


def main() -> None:
    write_documents(OUTPUT_DIR, DOCUMENTS)
    print(f"Generated {len(DOCUMENTS)} dummy documents in {OUTPUT_DIR}")

    write_documents(SIMILAR_COMPANY_OUTPUT_DIR, SIMILAR_COMPANY_DOCUMENTS)
    print(
        f"Generated {len(SIMILAR_COMPANY_DOCUMENTS)} conflict and irrelevant documents in {SIMILAR_COMPANY_OUTPUT_DIR}"
    )


if __name__ == "__main__":
    main()