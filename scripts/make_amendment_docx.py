from pathlib import Path

from docx import Document

workspace = Path(__file__).resolve().parent.parent
source_path = workspace / "demo" / "northwind-uplift-amendment-v2.txt"
output_path = workspace / "demo" / "northwind-uplift-amendment-v2.docx"

content = source_path.read_text(encoding="utf-8").split("\n\n")
document = Document()
for index, paragraph_text in enumerate(content):
    if index == 0:
        document.add_heading(paragraph_text.strip(), level=1)
    else:
        document.add_paragraph(paragraph_text.strip())

document.save(output_path)
print(output_path)
