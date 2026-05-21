"""Generate realistic contract documents for Pinnacle Logistics live demo."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "demo"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_clause(doc, number, title, text):
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {title}. ")
    run.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(10)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Document 1: Master Subscription Agreement
# ─────────────────────────────────────────────────────────────────────────────

def create_msa():
    doc = Document()

    add_heading(doc, "MASTER SUBSCRIPTION AGREEMENT")

    add_body(doc, (
        "This Master Subscription Agreement (\"Agreement\") is entered into as of "
        "April 1, 2025, by and between Conga Software, Inc., a Delaware corporation "
        "with its principal offices at 270 St. Paul Street, Denver, Colorado 80206 "
        "(\"Supplier\"), and Pinnacle Logistics Corp., a California corporation with "
        "its principal offices at 1400 Harbor Boulevard, Oakland, California 94607 "
        "(\"Customer\")."
    ))

    add_body(doc, (
        "WHEREAS, Customer desires to subscribe to Supplier's Conga Revenue "
        "Intelligence Suite (the \"Service\") for use by Customer's authorized users "
        "in connection with Customer's internal revenue operations, billing analytics, "
        "and contract lifecycle management activities; and"
    ))

    add_body(doc, (
        "WHEREAS, Supplier desires to provide such subscription subject to the terms "
        "and conditions set forth below;"
    ))

    add_body(doc, "NOW, THEREFORE, the parties agree as follows:")

    add_clause(doc, 1, "Definitions",
        "\"Authorized Users\" means Customer's employees and contractors who are "
        "permitted to access the Service under this Agreement. \"Subscription Term\" "
        "means the initial term and any renewal terms. \"Order Form\" means a "
        "mutually executed document specifying the Service edition, quantities, "
        "pricing, and any additional terms."
    )

    add_clause(doc, 2, "Subscription Term and Renewal",
        "The initial subscription term commences on April 1, 2025 and continues "
        "through March 31, 2026 (the \"Initial Term\"). Thereafter, the subscription "
        "will automatically renew for successive twelve (12) month renewal terms "
        "(each a \"Renewal Term\") unless either party provides written notice of "
        "non-renewal at least thirty (30) days prior to the end of the then-current term."
    )

    add_clause(doc, 3, "Fees and Invoicing",
        "Customer will be invoiced monthly for five hundred (500) subscribed platform "
        "seats at a rate of Two Hundred Dollars ($200.00) per seat per month, for a "
        "total monthly commitment of One Hundred Thousand Dollars ($100,000.00). "
        "Undisputed invoices are due and payable net thirty (30) days from the date "
        "of receipt."
    )

    add_clause(doc, 4, "Renewal Pricing Adjustment",
        "Beginning with the first Renewal Term and on each subsequent renewal "
        "anniversary thereafter, the recurring subscription fees shall be subject "
        "to a three percent (3%) annual price increase. Supplier may implement the "
        "3% annual price increase by providing Customer with at least thirty (30) "
        "days' prior written notice before the applicable renewal anniversary date."
    )

    add_clause(doc, 5, "Service Levels and Support",
        "Supplier will provide standard support during business hours (8:00 AM to "
        "6:00 PM Mountain Time, Monday through Friday, excluding US federal holidays). "
        "Supplier will use commercially reasonable efforts to maintain Service "
        "availability of at least 99.5% during each calendar month."
    )

    add_clause(doc, 6, "Data Security",
        "Supplier will maintain administrative, technical, and physical safeguards "
        "designed to protect Customer Data from unauthorized access, disclosure, "
        "alteration, or destruction. Supplier will comply with SOC 2 Type II "
        "requirements and will provide audit reports upon reasonable request."
    )

    add_clause(doc, 7, "Confidentiality",
        "Each party agrees to treat the other party's Confidential Information "
        "with at least the same degree of care it uses to protect its own "
        "confidential information of like kind, but in no event less than "
        "reasonable care."
    )

    add_clause(doc, 8, "Limitation of Liability",
        "EXCEPT FOR BREACHES OF SECTION 6 (DATA SECURITY) OR SECTION 7 "
        "(CONFIDENTIALITY), NEITHER PARTY'S AGGREGATE LIABILITY UNDER THIS "
        "AGREEMENT SHALL EXCEED THE TOTAL FEES PAID OR PAYABLE BY CUSTOMER "
        "DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY PRECEDING THE CLAIM."
    )

    add_clause(doc, 9, "Precedence",
        "Order forms, amendments, and renewal schedules executed after the "
        "effective date of this Agreement may supersede the commercial terms "
        "herein to the extent expressly stated therein. In the event of a "
        "conflict between this Agreement and an Order Form, the Order Form "
        "shall prevail solely with respect to pricing, quantities, and "
        "service-specific commercial terms."
    )

    add_clause(doc, 10, "General Provisions",
        "This Agreement, together with any Order Forms, amendments, and "
        "statements of work, constitutes the entire agreement between the "
        "parties with respect to the subject matter hereof and supersedes "
        "all prior or contemporaneous agreements, proposals, negotiations, "
        "and communications, whether written or oral."
    )

    add_body(doc, "")
    add_body(doc, "IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.")

    add_body(doc, "")
    add_body(doc, "CONGA SOFTWARE, INC.")
    add_body(doc, "By: /s/ Sarah Mitchell")
    add_body(doc, "Name: Sarah Mitchell")
    add_body(doc, "Title: VP, Commercial Operations")
    add_body(doc, "Date: March 28, 2025")

    add_body(doc, "")
    add_body(doc, "PINNACLE LOGISTICS CORP.")
    add_body(doc, "By: /s/ David Chen")
    add_body(doc, "Name: David Chen")
    add_body(doc, "Title: Chief Financial Officer")
    add_body(doc, "Date: March 29, 2025")

    path = OUTPUT_DIR / "pinnacle-master-subscription-agreement-v1.docx"
    doc.save(path)
    print(f"Created: {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# Document 2: Commercial Amendment (bumps uplift from 3% to 7%)
# ─────────────────────────────────────────────────────────────────────────────

def create_amendment():
    doc = Document()

    add_heading(doc, "COMMERCIAL AMENDMENT NO. 1")
    add_heading(doc, "to the Master Subscription Agreement dated April 1, 2025", level=2)

    add_body(doc, (
        "This Commercial Amendment No. 1 (\"Amendment\") is entered into as of "
        "January 15, 2026, by and between Conga Software, Inc. (\"Supplier\") and "
        "Pinnacle Logistics Corp. (\"Customer\"), and amends the Master Subscription "
        "Agreement dated April 1, 2025 (the \"Agreement\") solely with respect to "
        "the commercial pricing terms specified below."
    ))

    add_body(doc, "RECITALS")

    add_body(doc, (
        "A. The parties entered into the Agreement effective April 1, 2025, which "
        "currently provides for a 3% annual price increase at renewal."
    ))

    add_body(doc, (
        "B. Following a comprehensive account review and the addition of advanced "
        "analytics modules, predictive forecasting capabilities, and dedicated "
        "customer success resources, the parties have negotiated updated commercial "
        "terms that reflect the expanded scope of the engagement."
    ))

    add_body(doc, (
        "C. The parties desire to amend the Agreement to reflect the revised "
        "renewal pricing mechanics set forth herein."
    ))

    add_body(doc, "NOW, THEREFORE, the parties agree as follows:")

    add_clause(doc, 1, "Amendment to Section 4 (Renewal Pricing Adjustment)",
        "Section 4 of the Agreement is hereby deleted in its entirety and replaced "
        "with the following: \"Beginning with the Renewal Term commencing on April 1, "
        "2026, and on each subsequent renewal anniversary thereafter, the recurring "
        "subscription fees shall be subject to a seven percent (7%) annual price "
        "increase. Supplier may implement the 7% annual price increase by providing "
        "Customer with at least forty-five (45) days' prior written notice before "
        "the applicable renewal anniversary date.\""
    )

    add_clause(doc, 2, "Scope of Amendment",
        "This Amendment supersedes and replaces any prior pricing adjustment "
        "provisions in the Agreement or any earlier Order Form that conflict with "
        "the terms set forth in Section 1 above. All other terms and conditions of "
        "the Agreement remain in full force and effect."
    )

    add_clause(doc, 3, "Effective Date",
        "This Amendment shall become effective upon the last signature below and "
        "shall apply to all renewal periods beginning on or after April 1, 2026."
    )

    add_clause(doc, 4, "Consideration",
        "In consideration of the revised pricing mechanics, Supplier agrees to "
        "provide Customer with access to the Advanced Analytics Add-On module and "
        "a dedicated Customer Success Manager at no additional charge for the "
        "duration of the current Renewal Term."
    )

    add_clause(doc, 5, "Counterparts",
        "This Amendment may be executed in counterparts, each of which shall be "
        "deemed an original, and all of which together shall constitute one and "
        "the same instrument."
    )

    add_body(doc, "")
    add_body(doc, "IN WITNESS WHEREOF, the parties have executed this Amendment as of the date first written above.")

    add_body(doc, "")
    add_body(doc, "CONGA SOFTWARE, INC.")
    add_body(doc, "By: /s/ Sarah Mitchell")
    add_body(doc, "Name: Sarah Mitchell")
    add_body(doc, "Title: VP, Commercial Operations")
    add_body(doc, "Date: January 15, 2026")

    add_body(doc, "")
    add_body(doc, "PINNACLE LOGISTICS CORP.")
    add_body(doc, "By: /s/ David Chen")
    add_body(doc, "Name: David Chen")
    add_body(doc, "Title: Chief Financial Officer")
    add_body(doc, "Date: January 14, 2026")

    path = OUTPUT_DIR / "pinnacle-commercial-amendment-v1.docx"
    doc.save(path)
    print(f"Created: {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# Document 3: Renewal Notice (confirms the 7% and sends notice)
# ─────────────────────────────────────────────────────────────────────────────

def create_renewal_notice():
    doc = Document()

    add_heading(doc, "RENEWAL PRICING NOTICE")

    add_body(doc, "Date: February 12, 2026")
    add_body(doc, "To: Pinnacle Logistics Corp.")
    add_body(doc, "Attn: David Chen, Chief Financial Officer")
    add_body(doc, "From: Conga Software, Inc., Revenue Operations")
    add_body(doc, "Re: Annual Renewal Pricing Adjustment — Conga Revenue Intelligence Suite")
    add_body(doc, "Contract Reference: Master Subscription Agreement dated April 1, 2025, as amended by Commercial Amendment No. 1 dated January 15, 2026")

    add_body(doc, "")
    add_body(doc, "Dear Mr. Chen,")

    add_body(doc, (
        "Pursuant to Commercial Amendment No. 1 to the Master Subscription "
        "Agreement (collectively, the \"Agreement\"), this letter serves as "
        "formal notice that the annual renewal pricing adjustment will take "
        "effect for the Renewal Term commencing April 1, 2026."
    ))

    add_heading(doc, "Renewal Pricing Details", level=2)

    add_body(doc, (
        "In accordance with the amended Section 4 of the Agreement, the "
        "recurring subscription fees for Conga Revenue Intelligence Suite "
        "will be increased by seven percent (7%) effective April 1, 2026. "
        "The updated monthly fee will be calculated as follows:"
    ))

    add_body(doc, "• Current monthly fee: $100,000.00 (500 seats × $200.00/seat)")
    add_body(doc, "• Annual uplift percentage: 7%")
    add_body(doc, "• New monthly fee effective April 1, 2026: $107,000.00 (500 seats × $214.00/seat)")

    add_body(doc, (
        "This notice is being provided more than forty-five (45) days in advance "
        "of the April 1, 2026 renewal anniversary, as required by the Agreement. "
        "The notice deadline per the amendment was February 14, 2026; this notice "
        "is delivered on February 12, 2026."
    ))

    add_heading(doc, "Continuation of Service", level=2)

    add_body(doc, (
        "Unless you provide written notice of non-renewal at least thirty (30) "
        "days before March 31, 2026, the subscription will automatically renew "
        "for an additional twelve (12) month term at the adjusted pricing."
    ))

    add_body(doc, (
        "We value our partnership with Pinnacle Logistics and look forward to "
        "continuing to support your revenue operations initiatives. Please do "
        "not hesitate to reach out with any questions."
    ))

    add_body(doc, "")
    add_body(doc, "Sincerely,")
    add_body(doc, "")
    add_body(doc, "Sarah Mitchell")
    add_body(doc, "VP, Commercial Operations")
    add_body(doc, "Conga Software, Inc.")
    add_body(doc, "sarah.mitchell@conga.com")

    add_body(doc, "")
    add_body(doc, "cc: Pinnacle Logistics Account Team")
    add_body(doc, "    Conga Revenue Operations")

    path = OUTPUT_DIR / "pinnacle-renewal-notice-2026.docx"
    doc.save(path)
    print(f"Created: {path}")
    return path


if __name__ == "__main__":
    create_msa()
    create_amendment()
    create_renewal_notice()
    print("\nAll documents created. Ready for upload via the UI.")
