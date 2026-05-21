from __future__ import annotations

from pathlib import Path

import psycopg
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.config import get_data_settings, get_object_store_settings
from app.object_store import get_minio_client


OUTPUT_DIR = Path(__file__).resolve().parent.parent / "data" / "generated_documents"

DOCUMENT_SECTION_OVERRIDES: dict[str, list[tuple[str, list[str]]]] = {
    "doc-1101": [
        (
            "Pricing Override",
            [
                "For commercial pricing only, this order form supersedes the master agreement and sets the first renewal uplift at 6% of recurring subscription charges.",
                "Supplier may implement the 6% annual uplift for the next renewal term by providing Customer at least 30 days notice before the applicable renewal anniversary.",
            ],
        ),
        (
            "Commercial Notes",
            [
                "All other terms of the master subscription agreement remain unchanged unless expressly modified in a later amendment.",
            ],
        ),
    ],
    "doc-1102": [
        (
            "Commercial Amendment",
            [
                "Effective for the 2026 renewal planning cycle, the parties amend the renewal pricing clause to permit a 7% annual uplift.",
                "Supplier may apply the 7% annual uplift upon providing Customer at least 45 days notice before the relevant renewal date.",
            ],
        ),
        (
            "Precedence",
            [
                "This amendment controls over any inconsistent pricing terms in the original agreement or order form until a later amendment is executed.",
            ],
        ),
    ],
    "doc-1103": [
        (
            "Commercial Amendment",
            [
                "The parties agree that, beginning with the next renewal term, subscription fees are subject to a 9% annual price increase.",
                "Supplier may implement the 9% annual price increase by providing Customer at least 60 days notice before the applicable renewal anniversary.",
            ],
        ),
        (
            "Precedence",
            [
                "This amendment supersedes prior 5%, 6%, and 7% pricing references solely with respect to renewal uplift mechanics.",
            ],
        ),
    ],
    "doc-1104": [
        (
            "Renewal Pricing Memo",
            [
                "The renewal team prepared a customer-facing pricing memo confirming the commercial amendment target of a 9% annual uplift for the 2026 renewal cycle.",
                "The memo states that 60 days notice is required before the renewal anniversary for the 9% uplift to take effect.",
            ],
        )
    ],
    "doc-1201": [
        (
            "Pricing Amendment",
            [
                "This pricing amendment updates the Apex commercial package to a 4% annual uplift for the next renewal term.",
                "The 4% uplift may be implemented if Supplier delivers no less than 21 days notice before renewal.",
            ],
        )
    ],
    "doc-1202": [
        (
            "Renewal Schedule",
            [
                "The most recent renewal schedule supersedes prior pricing amendments and sets the 2026 renewal uplift at 6% of recurring subscription fees.",
                "Supplier must send the formal renewal notice at least 14 days before June 1, 2026 in order to implement the 6% uplift.",
            ],
        )
    ],
    "doc-1203": [
        (
            "Draft Renewal Notice",
            [
                "This draft renewal notice references a 6% annual uplift and a final outbound notice deadline of May 18, 2026.",
                "The document was prepared internally and does not itself confirm that notice has been sent.",
            ],
        )
    ],
    "doc-1402": [
        (
            "Order Form Pricing",
            [
                "For Conga Quote to Cash Advanced, the executed order form overrides the master agreement pricing mechanics and sets a 6% annual uplift at renewal.",
                "Supplier may apply the 6% uplift by providing Customer at least 30 days notice before the renewal anniversary.",
            ],
        )
    ],
    "doc-1403": [
        (
            "Commercial Amendment",
            [
                "The parties later executed a commercial amendment providing that recurring subscription charges are subject to an 8% annual uplift beginning January 1, 2026.",
                "Supplier may implement the 8% annual uplift upon providing Customer at least 60 days notice before the applicable renewal anniversary.",
            ],
        ),
        (
            "Priority of Terms",
            [
                "This amendment supersedes the earlier 4% master agreement language and 6% order form pricing solely for renewal uplift calculations.",
            ],
        ),
    ],
    "doc-1404": [
        (
            "Renewal Operations Brief",
            [
                "Revenue operations documented the 2026 renewal assumption as an 8% annual uplift with a 60 day notice requirement.",
                "The brief references the controlling commercial amendment as the source of record for pricing changes.",
            ],
        )
    ],
    "doc-1502": [
        (
            "Renewal Schedule",
            [
                "The Redwood renewal schedule increases the upcoming renewal uplift from the base agreement level to 4% of recurring subscription fees.",
                "Supplier must provide at least 21 days notice before the renewal date to apply the 4% uplift.",
            ],
        )
    ],
    "doc-1503": [
        (
            "Commercial Amendment",
            [
                "The later Redwood commercial amendment sets the next renewal uplift at 5% and supersedes any inconsistent 2% or 4% pricing references.",
                "Supplier may implement the 5% uplift with at least 30 days notice before June 25, 2026.",
            ],
        )
    ],
    "doc-1504": [
        (
            "Renewal Playbook",
            [
                "The renewal playbook confirms the operational target of a 5% annual uplift and identifies May 26, 2026 as the outbound notice deadline.",
                "The document is an internal planning artifact and does not confirm that notice has yet been sent.",
            ],
        )
    ],
}


def _query_document_rows() -> list[dict[str, object]]:
    settings = get_data_settings()
    if not settings.database_url:
        raise RuntimeError("DATABASE_URL must be configured to seed documents")

    query = """
        SELECT
            document_id,
            contract_documents.contract_id,
            document_type,
            file_name,
            mime_type,
            storage_key,
            version,
            page_count,
            accounts.name AS account_name,
            product_name,
            term_start,
            term_end,
            base_price,
            currency,
            quantity,
            raw_contract_text
        FROM contract_documents
        JOIN contracts ON contracts.contract_id = contract_documents.contract_id
        JOIN accounts ON accounts.account_id = contracts.account_id
        WHERE contract_documents.document_id NOT LIKE 'doc-upload-%'
        ORDER BY contract_documents.document_id
    """

    with psycopg.connect(settings.database_url) as connection:
        with connection.cursor(row_factory=psycopg.rows.dict_row) as cursor:
            cursor.execute(query)
            return list(cursor.fetchall())


def _build_common_sections(row: dict[str, object]) -> list[tuple[str, list[str]]]:
    return [
        (
            "Background",
            [
                (
                    f"This agreement is entered into between Conga Software, Inc. and {row['account_name']} for the subscription, support, and use of {row['product_name']}. "
                    "The parties intend this document to operate as a real commercial instrument governing entitlements, pricing, invoicing, support obligations, confidentiality, and termination."
                ),
                (
                    "Customer's use of the services is limited to internal business purposes, subject to the ordering documents, acceptable use policy, and any security or compliance schedules incorporated by reference."
                ),
            ],
        ),
        (
            "Commercial Terms",
            [
                (
                    f"The initial committed term begins on {row['term_start']:%B %d, %Y} and continues through {row['term_end']:%B %d, %Y}. "
                    f"Customer is subscribing to {row['quantity']} units at a base price of {row['currency']} {row['base_price']:.2f} per unit, billed in accordance with the invoice cadence defined in the applicable order form."
                ),
                (
                    "Unless otherwise stated, undisputed invoices are due net thirty (30) days from receipt. Late payments may accrue interest at the lesser of one and one-half percent per month or the maximum amount permitted by law."
                ),
            ],
        ),
        (
            "Security, Compliance, and Operations",
            [
                (
                    "Supplier will maintain administrative, technical, and physical safeguards designed to protect Customer data against unauthorized access, disclosure, or loss. Supplier will provide standard support during normal business hours and commercially reasonable efforts to maintain service availability in accordance with the support policy."
                ),
                (
                    "Customer remains responsible for user access management, lawful instructions concerning Customer data, and compliance with industry-specific obligations applicable to Customer's regulated operations."
                ),
            ],
        ),
        (
            "Legal Terms",
            [
                (
                    "This agreement, together with each incorporated order form, amendment, and policy attachment, constitutes the entire agreement between the parties with respect to the subject matter hereof and supersedes prior or contemporaneous proposals and communications."
                ),
                (
                    "Neither party may assign this agreement without the prior written consent of the other party, except in connection with a merger, acquisition, or sale of substantially all assets. Governing law and venue shall be identified in the final signature version maintained by Supplier."
                ),
            ],
        ),
    ]


def _document_specific_sections(row: dict[str, object]) -> list[tuple[str, list[str]]]:
    override_sections = DOCUMENT_SECTION_OVERRIDES.get(str(row["document_id"]))
    if override_sections is not None:
        return override_sections

    if row["document_type"] == "order_form":
        return [
            (
                "Ordering Schedule",
                [
                    f"Product: {row['product_name']}",
                    f"Subscribed Quantity: {row['quantity']} seats or units",
                    f"Commercial Start Date: {row['term_start']:%B %d, %Y}",
                    f"Committed End Date: {row['term_end']:%B %d, %Y}",
                ],
            ),
            (
                "Renewal and Price Adjustment",
                [
                    "Upon each renewal term, recurring subscription charges may be adjusted in accordance with the governing agreement and any applicable notice requirements set out below.",
                    str(row["raw_contract_text"]),
                ],
            ),
        ]

    return [
        (
            "Renewal Pricing Adjustment",
            [
                str(row["raw_contract_text"]),
                (
                    "If an amendment or order form expressly overrides the renewal pricing mechanics described in this agreement, the later-dated document controls solely to the extent of the inconsistency."
                ),
            ],
        ),
        (
            "Signature Page",
            [
                f"Accepted and agreed by authorized representatives of Conga Software, Inc. and {row['account_name']}.",
                "Name: ____________________    Title: ____________________    Date: ____________________",
                "Name: ____________________    Title: ____________________    Date: ____________________",
            ],
        ),
    ]


def _build_pdf_document(row: dict[str, object], destination: Path) -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocumentTitle",
        parent=styles["Title"],
        fontSize=22,
        leading=28,
        alignment=TA_CENTER,
        spaceAfter=18,
    )
    heading_style = ParagraphStyle("SectionHeading", parent=styles["Heading2"], fontSize=14, leading=18, spaceAfter=10)
    body_style = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10.5, leading=15, spaceAfter=10)
    small_style = ParagraphStyle("Small", parent=styles["BodyText"], fontSize=9.5, leading=13, spaceAfter=8, textColor=colors.HexColor("#555555"))

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(str(destination), pagesize=LETTER, leftMargin=56, rightMargin=56, topMargin=54, bottomMargin=50)
    story = [
        Paragraph(str(row["file_name"]).replace("-", " ").replace(".pdf", "").upper(), title_style),
        Paragraph(f"Customer: {row['account_name']}", small_style),
        Paragraph(f"Product: {row['product_name']}", small_style),
        Spacer(1, 12),
    ]

    if row["document_type"] == "order_form":
        order_table = Table(
            [
                ["Commercial Item", "Value"],
                ["Product", str(row["product_name"])],
                ["Subscribed Quantity", str(row["quantity"])],
                ["Base Unit Price", f"{row['currency']} {float(row['base_price']):.2f}"],
                ["Committed Term", f"{row['term_start']:%b %d, %Y} to {row['term_end']:%b %d, %Y}"],
            ],
            colWidths=[160, 300],
        )
        order_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e9dfcc")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1f1b16")),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#b9a88c")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.white),
                ]
            )
        )
        story.extend([order_table, Spacer(1, 18)])

    sections = _build_common_sections(row) + _document_specific_sections(row)
    for index, (heading, paragraphs) in enumerate(sections):
        story.append(Paragraph(heading, heading_style))
        for paragraph in paragraphs:
            story.append(Paragraph(paragraph.replace("\n", "<br/>"), body_style))
        if index in {1, 3}:
            story.append(PageBreak())

    def draw_page(canvas, _document):
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(colors.HexColor("#555555"))
        canvas.drawRightString(560, 30, f"Page {canvas.getPageNumber()}")
        canvas.drawString(56, 30, "Confidential - synthetic enterprise document")
        canvas.restoreState()

    document.build(story, onFirstPage=draw_page, onLaterPages=draw_page)


def _build_docx_document(row: dict[str, object], destination: Path) -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(str(row["file_name"]).replace("-", " ").replace(".docx", "").upper())
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph(f"Customer: {row['account_name']}\nProduct: {row['product_name']}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    common_sections = _build_common_sections(row)
    for page_index, section_group in enumerate((common_sections[:2], common_sections[2:] + _document_specific_sections(row))):
        for heading, paragraphs in section_group:
            document.add_heading(heading, level=2)
            for paragraph in paragraphs:
                document.add_paragraph(paragraph)
        if page_index == 0:
            document.add_page_break()

    annex = document.add_section(WD_SECTION_START.NEW_PAGE)
    annex.top_margin = Inches(0.7)
    annex.bottom_margin = Inches(0.6)
    annex.left_margin = Inches(0.8)
    annex.right_margin = Inches(0.8)
    document.add_heading("Schedule A - Operational Assumptions", level=2)
    document.add_paragraph(
        "For clarity, pricing operations, entitlement changes, and renewal approvals shall be administered in accordance with Supplier's standard commercial workflow, with all exceptions documented through a written amendment or updated order form."
    )
    document.add_paragraph(
        "This schedule is included to make the seeded document resemble a real customer-ready commercial artifact rather than a single-clause excerpt."
    )

    document.save(destination)


def _generate_document(row: dict[str, object]) -> Path:
    destination = OUTPUT_DIR / str(row["file_name"])
    if row["mime_type"] == "application/pdf":
        _build_pdf_document(row, destination)
    else:
        _build_docx_document(row, destination)
    return destination


def _upload_and_mark(row: dict[str, object], file_path: Path) -> None:
    object_settings = get_object_store_settings()
    bucket_name = object_settings.bucket_name or "contract-documents"
    client = get_minio_client()
    client.fput_object(bucket_name, str(row["storage_key"]), str(file_path), content_type=str(row["mime_type"]))

    settings = get_data_settings()
    with psycopg.connect(settings.database_url) as connection:
        with connection.cursor() as cursor:
            cursor.execute(
                (
                    "UPDATE contract_documents "
                    "SET ingestion_status = CASE "
                    "WHEN EXISTS (SELECT 1 FROM obligation_extractions WHERE document_id = %s) THEN 'parsed' "
                    "ELSE 'uploaded' END "
                    "WHERE document_id = %s"
                ),
                (str(row["document_id"]), str(row["document_id"])),
            )
        connection.commit()


def main() -> None:
    rows = _query_document_rows()
    for row in rows:
        file_path = _generate_document(row)
        _upload_and_mark(row, file_path)
        print(f"Seeded {row['document_id']} -> {file_path.name}")


if __name__ == "__main__":
    main()