from __future__ import annotations

from io import BytesIO
from re import Pattern

from docx import Document as DocxDocument
from pypdf import PdfReader


MIME_BY_SUFFIX = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _extract_docx_text(payload: bytes) -> tuple[str, int | None, int | None]:
    document = DocxDocument(BytesIO(payload))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells = []
    for table in document.tables:
        for row in table.rows:
            cell_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if cell_text:
                table_cells.append(cell_text)
    text = "\n\n".join(paragraphs + table_cells).strip()
    return text, None, None


def _extract_pdf_text(
    payload: bytes,
    *,
    page_hint_pattern: Pattern[str] | None = None,
) -> tuple[str, int | None, int | None]:
    reader = PdfReader(BytesIO(payload))
    page_texts: list[str] = []
    matched_page: int | None = None

    for page_number, page in enumerate(reader.pages, start=1):
        extracted_text = (page.extract_text() or "").strip()
        if extracted_text:
            page_texts.append(extracted_text)
            if matched_page is None and page_hint_pattern and page_hint_pattern.search(extracted_text):
                matched_page = page_number

    return "\n\n".join(page_texts).strip(), len(reader.pages), matched_page


def extract_document_text(
    mime_type: str,
    payload: bytes,
    *,
    page_hint_pattern: Pattern[str] | None = None,
) -> tuple[str, int | None, int | None, str]:
    if mime_type == MIME_BY_SUFFIX[".pdf"]:
        text, page_count, page_number = _extract_pdf_text(payload, page_hint_pattern=page_hint_pattern)
        return text, page_count, page_number, "pdf-native-text"

    text, page_count, page_number = _extract_docx_text(payload)
    return text, page_count, page_number, "docx-native-text"